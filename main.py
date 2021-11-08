import sys

from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog, QInputDialog, QFileDialog, QComboBox
import sqlite3
from PyQt5.QtCore import Qt
from docx import Document
import random as r


class MainMenuWidget(QMainWindow):
    def __init__(self, parent=None):
        super(MainMenuWidget, self).__init__(parent)
        uic.loadUi('Main_Menu.ui', self)
        self.pushButton_2.clicked.connect(self.open_base_data)
        self.pushButton.clicked.connect(self.open_generate_work)
        self.pushButton_3.clicked.connect(self.exit)

    def open_base_data(self):
        f1 = BaseData(self)
        f1.show()
        self.hide()

    def open_generate_work(self):
        f = CreateWork(self)
        f.show()
        self.hide()

    def exit(self):
        self.close()


class BaseData(QMainWindow):
    def __init__(self, parent=None):
        super(BaseData, self).__init__(parent)
        uic.loadUi('Base_Data.ui', self)
        self.pushButton_2.clicked.connect(self.returned)
        self.pushButton.clicked.connect(self.ex4)
        self.pushButton_3.clicked.connect(self.add_theme)
        self.comboBox.hide()

    def ex4(self):
        self.con = sqlite3.connect("PhB.db")
        cur = self.con.cursor()
        a = cur.execute("select name from sqlite_master where type = 'table' and name not like '_all'").fetchall()
        self.comboBox.show()
        for i in a:
            self.comboBox.addItem(str(i[0]))
        self.comboBox.show()
        self.comboBox.activated[str].connect(self.onChanged)
        self.con.close()

    def onChanged(self, text):
        self.hide()
        a = WriteQue(self, text)
        a.show()

    def returned(self):
        self.hide()
        ex1 = MainMenuWidget(self)
        ex1.show()

    def add_theme(self):
        text, ok = QInputDialog.getText(self, 'Добавить раздел',
                                        'Введите название')
        if text:
            if ok:
                con = sqlite3.connect("PhB.db")
                cur = con.cursor()
                a = "CREATE TABLE " + '"' + text + '"' + " (body STRING NOT NULL,list STRING NOT NULL,formula STRING NOT NULL);"
                cur.execute(a)
                con.commit()
                con.close()


class WriteQue(QMainWindow):
    def __init__(self, parent=None, text=None):
        super(QMainWindow, self).__init__(parent)
        uic.loadUi('Write_Que.ui', self)
        self.text = text
        self.pushButton.hide()
        self.pushButton.clicked.connect(self.save_result)
        self.clear()
        self.pushButton_2.clicked.connect(self.ready)
        self.spinBox.setMaximum(100000000)
        self.spinBox_2.setMaximum(100000000)
        self.spinBox_3.setMaximum(100000000)
        self.spinBox_4.setMaximum(100000000)
        self.spinBox_5.setMaximum(100000000)
        self.spinBox_6.setMaximum(100000000)
        self.spinBox_7.setMaximum(100000000)
        self.spinBox_8.setMaximum(100000000)
        self.spinBox_9.setMaximum(100000000)
        self.spinBox_10.setMaximum(100000000)
        self.spinBox.setMinimum(-100000000)
        self.spinBox_2.setMinimum(-100000000)
        self.spinBox_3.setMinimum(-100000000)
        self.spinBox_4.setMinimum(-100000000)
        self.spinBox_5.setMinimum(-100000000)
        self.spinBox_6.setMinimum(-100000000)
        self.spinBox_7.setMinimum(-100000000)
        self.spinBox_8.setMinimum(-100000000)
        self.spinBox_9.setMinimum(-100000000)
        self.spinBox_10.setMinimum(-100000000)
        self.label_404.hide()
        self.lineEdit.hide()
        self.labelblya.hide()

    def save_result(self):
        con = sqlite3.connect("PhB.db")
        cur = con.cursor()
        list_of_spin_bx = str(self.spinBox.value()) + ' ' + str(self.spinBox_2.value()) + ' '\
                          + str(self.spinBox_3.value()) + ' ' + str(self.spinBox_4.value()) + ' '\
                          + str(self.spinBox_5.value()) + ' ' + str(self.spinBox_6.value()) + ' '\
                          + str(self.spinBox_7.value()) + ' ' + str(self.spinBox_8.value()) + ' '\
                          + str(self.spinBox_9.value()) + ' ' + str(self.spinBox_10.value())
        b = self.lineEdit.text()
        a = "INSERT INTO " + '"' + self.text + '"' + "(body, list, formula) VALUES(" + '"'\
            + self.plainTextEdit.toPlainText() + '", "(' + list_of_spin_bx + ')", ' + '"' + str(b) + '"' + ')'
        b = "INSERT INTO _all(body, list, formula) VALUES(" + '"'\
            + self.plainTextEdit.toPlainText() + '", "(' + list_of_spin_bx + ')", ' + '"' + str(b) + '"' + ')'
        cur.execute(a)
        cur.execute(b)
        con.commit()
        con.close()
        self.close()
        a = MainMenuWidget(self)
        a.show()

    def show_none_data_dialog(self):
        self.hide()
        a = NoneDataDialog(self)
        a.show()

    def activ_1(self):
        self.clear()
        self.spinBox.show()
        self.spinBox_2.show()
        self.label_1.show()
        self.label_2.show()
        self.label_3.show()

    def activ_2(self):
        self.clear()
        self.activ_1()
        self.spinBox_3.show()
        self.spinBox_4.show()
        self.label_4.show()
        self.label_5.show()
        self.label_6.show()

    def activ_3(self):
        self.clear()
        self.activ_2()
        self.spinBox_5.show()
        self.spinBox_6.show()
        self.label_7.show()
        self.label_8.show()
        self.label_9.show()

    def activ_4(self):
        self.clear()
        self.activ_3()
        self.spinBox_7.show()
        self.spinBox_8.show()
        self.label_7.show()
        self.label_8.show()
        self.label_10.show()
        self.label_11.show()
        self.label_12.show()

    def activ_5(self):
        self.clear()
        self.activ_4()
        self.spinBox_9.show()
        self.spinBox_10.show()
        self.label_13.show()
        self.label_14.show()
        self.label_15.show()

    def clear(self):
        self.spinBox.hide()
        self.spinBox_2.hide()
        self.spinBox_3.hide()
        self.spinBox_4.hide()
        self.spinBox_5.hide()
        self.spinBox_6.hide()
        self.spinBox_7.hide()
        self.spinBox_8.hide()
        self.spinBox_9.hide()
        self.spinBox_10.hide()
        self.label_1.hide()
        self.label_2.hide()
        self.label_3.hide()
        self.label_4.hide()
        self.label_5.hide()
        self.label_6.hide()
        self.label_7.hide()
        self.label_8.hide()
        self.label_9.hide()
        self.label_10.hide()
        self.label_11.hide()
        self.label_12.hide()
        self.label_13.hide()
        self.label_14.hide()
        self.label_15.hide()

    def ready(self):
        text = self.plainTextEdit.toPlainText()
        if not text:
            a = NoneDataDialog(self)
            a.show()
        elif text.count('!') > 5:
            a = ManyDataDialog(self)
            a.show()
        elif True:
            if text.count('!') == 5:
                self.activ_5()
            elif text.count('!') == 4:
                self.activ_4()
            elif text.count('!') == 3:
                self.activ_3()
            elif text.count('!') == 2:
                self.activ_2()
            elif text.count('!') == 1:
                self.activ_1()
            elif text.count('!') == 0:
                self.clear()
            self.labelblya.setText(str(text))
            self.labelblya.show()
            self.pushButton.show()
            self.plainTextEdit.hide()
            self.pushButton_2.hide()
            self.label_404.show()
            self.lineEdit.show()


class CreateWork(QMainWindow):
    def __init__(self, parent=None):
        super(CreateWork, self).__init__(parent)
        uic.loadUi('Generate_Work.ui', self)
        self.checklist = [False, False, False]
        self.pushButton.clicked.connect(self.ready)
        self.checkBox.stateChanged.connect(self.ch1)
        self.checkBox_2.stateChanged.connect(self.ch2)
        self.checkBox_3.stateChanged.connect(self.ch3)

    def ch1(self, state):
        if state == Qt.Checked:
            self.checklist = [True, self.checklist[1], self.checklist[2]]
        else:
            self.checklist = [False, self.checklist[1], self.checklist[2]]

    def ch2(self, state):
        if state == Qt.Checked:
            self.checklist = [self.checklist[0], True, self.checklist[2]]
        else:
            self.checklist = [self.checklist[0], False, self.checklist[2]]

    def ch3(self, state):
        if state == Qt.Checked:
            self.checklist = [self.checklist[0], self.checklist[1], True]
        else:
            self.checklist = [self.checklist[0], self.checklist[1], False]

    def ready(self):
        a = SolveWork(self, self.checklist, self.spinBox.value())
        a.show()
        self.hide()


class NoneDataDialog(QDialog):
    def __init__(self, parent=None):
        super(NoneDataDialog, self).__init__(parent)
        uic.loadUi('None_Data_Dialog.ui', self)
        self.pushButton.clicked.connect(self.close)

    def close(self):
        self.hide()


class ManyDataDialog(QDialog):
    def __init__(self, parent=None):
        super(ManyDataDialog, self).__init__(parent)
        uic.loadUi('Many_Data_Dialog.ui', self)
        self.pushButton.clicked.connect(self.close)

    def close(self):
        self.hide()


class SolveWork(QMainWindow):
    def __init__(self, parent=None, tlist=None, count=None):
        super(SolveWork, self).__init__(parent)
        uic.loadUi('Solve_Cases.ui', self)
        self.otlad = tlist
        self.unical = tlist[0]
        self.new_list = tlist[1]
        self.answers = tlist[2]
        self.count = count
        self.pushButton_3.clicked.connect(self.close)
        self.pushButton.clicked.connect(self.add_case)
        self.list_of_cases = []
        self.pushButton_4.clicked.connect(self.choose_directory)
        self.pushButton_6.hide()
        self.pushButton_6.clicked.connect(self.add_theme)

    def close(self):
        self.hide()
        a = CreateWork(self)
        a.show()

    def add_case(self):
        self.a11 = QComboBox(self)
        self.a11.resize(100, 30)
        self.a11.move(200, 110)
        self.a11.show()
        self.con = sqlite3.connect("PhB.db")
        cur = self.con.cursor()
        a = cur.execute("select name from sqlite_master where type = 'table' and name not like '_all'").fetchall()
        for i in a:
            self.a11.addItem(str(i[0]))
        self.con.close()
        self.a11.activated[str].connect(self.on_changed)
        self.con.close()

    def on_changed(self, text):
        self.b11 = QComboBox(self)
        self.b11.resize(100, 30)
        self.b11.move(340, 110)
        self.b11.show()
        self.con = sqlite3.connect("PhB.db")
        cur = self.con.cursor()
        a = cur.execute("select body from" + "'" + text + "'").fetchall()
        for i in a:
            self.b11.addItem(str(i[0]))
        self.b11.activated[str].connect(self.case_changed)
        self.con.close()

    def case_changed(self, text):
        self.list_of_cases.append(text)
        self.listWidget.clear()
        self.listWidget.addItems(self.list_of_cases)
        self.a11.hide()
        self.b11.hide()

    def choose_directory(self):
        dirlist = QFileDialog.getExistingDirectory(self, "Выбрать папку", ".")
        self.path = dirlist
        self.label.setText(str(dirlist))
        self.pushButton_6.show()

    def cook_case(self, name):
        self.con = sqlite3.connect("PhB.db")
        cur = self.con.cursor()
        c = cur.execute("select * from _all where body like " + '"' + name + '"').fetchall()[0]
        text, data, formula = c[0], c[1], c[2]
        data = data[1:-1]
        data = data.split(' ')
        data = list(map(int, data))
        a1, a2 = data[0], data[1]
        b1, b2 = data[2], data[3]
        c1, c2 = data[4], data[5]
        d1, d2 = data[6], data[7]
        e1, e2 = data[8], data[9]
        a = r.randint(a1, a2)
        b = r.randint(b1, b2)
        c = r.randint(c1, c2)
        d = r.randint(d1, d2)
        e = r.randint(e1, e2)
        text = text.replace('1(!)', str(a))
        text = text.replace('2(!)', str(b))
        text = text.replace('3(!)', str(c))
        text = text.replace('4(!)', str(d))
        text = text.replace('5(!)', str(e))
        formula = formula.replace('1(!)', str(a))
        formula = formula.replace('2(!)', str(b))
        formula = formula.replace('3(!)', str(c))
        formula = formula.replace('4(!)', str(d))
        formula = formula.replace('5(!)', str(e))
        ans = eval(formula)
        self.con.close()
        return text, str(ans)

    def work_with_word(self, name):
        document = Document()
        slov_of_ans = {}
        if self.unical:
            for i in range(1, self.count + 1):
                document.add_paragraph(("Вариант " + str(i)))
                for j in self.list_of_cases:
                    text, ans = self.cook_case(j)
                    document.add_paragraph(text)
                    if i in slov_of_ans.keys():
                        a = slov_of_ans[i]
                        a.append(ans)
                        slov_of_ans[i] = a
                    else:
                        slov_of_ans[i] = [ans]
                if self.new_list:
                    document.add_page_break()
            if self.answers:
                for i in slov_of_ans.keys():
                    document.add_paragraph(str(i))
                    document.add_paragraph(str(', '.join(slov_of_ans[i])))
        else:
            ansers = []
            texts = []
            for j in self.list_of_cases:
                text, ans = self.cook_case(j)
                texts.append(text)
                ansers.append(ans)
            for i in range(1, self.count + 1):
                document.add_paragraph(("Вариант" + str(i)))
                for k in texts:
                    document.add_paragraph(k)
                if self.new_list:
                    document.add_page_break()
            if self.answers:
                document.add_page_break()
                document.add_paragraph(' '.join(ansers))
        document.save(self.path + '\\' + name + '.docx')

    def add_theme(self):
        text, ok = QInputDialog.getText(self, 'Сохранить как (название файла)',
                                        'Введите название файла')
        if text:
            if ok:
                self.work_with_word(text)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MainMenuWidget()
    ex.show()
    sys.exit(app.exec_())
